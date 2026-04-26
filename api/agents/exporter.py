"""ExperimentPlan -> Markdown -> {pdf, docx, tex, md}.

PDF uses xhtml2pdf (pure-Python HTML->PDF, no system deps required).
DOCX and LaTeX use pandoc when available.
Plain markdown is always available.
"""
from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any


def _pandoc() -> str | None:
    found = shutil.which("pandoc")
    if found:
        return found
    candidates = [
        Path(os.environ.get("LOCALAPPDATA", "")) / "Pandoc" / "pandoc.exe",
        Path(os.environ.get("ProgramFiles", "")) / "Pandoc" / "pandoc.exe",
    ]
    for c in candidates:
        if c.exists():
            return str(c)
    return None


# ─── Markdown render ──────────────────────────────────────────────────────
def plan_to_markdown(plan: dict[str, Any]) -> str:
    lines: list[str] = []
    p = lambda s="": lines.append(s)

    p(f"# {plan.get('title', 'Experiment Plan')}")
    p()
    p(f"> *{plan.get('hypothesis', '')}*")
    p()
    p("## Novelty")
    p(plan.get("novelty_summary", ""))
    p()

    env = plan.get("environmental_conditions") or {}
    p("## Environmental Conditions")
    p(f"- **Temperature:** {env.get('temp_min_C')}–{env.get('temp_max_C')} °C")
    if env.get("humidity_min_pct") is not None:
        p(f"- **Humidity:** {env.get('humidity_min_pct')}–{env.get('humidity_max_pct')} %")
    if env.get("light"):
        p(f"- **Light:** {env['light']}")
    if env.get("atmosphere"):
        p(f"- **Atmosphere:** {env['atmosphere']}")
    if env.get("season_sensitivity"):
        p(f"- **Season sensitivity:** {env['season_sensitivity']}")
    p()

    p("## Protocol")
    for i, s in enumerate(plan.get("protocol") or [], 1):
        parallel = ", ".join(s.get("can_run_parallel_with") or [])
        p(f"### Step {i} — {s.get('name')}  *({s.get('duration_minutes', 0)} min)*")
        p(f"**Rationale.** {s.get('rationale', '')}")
        if s.get("materials_used"):
            p(f"- Materials: {', '.join(s['materials_used'])}")
        if s.get("equipment_used"):
            p(f"- Equipment: {', '.join(s['equipment_used'])}")
        if s.get("qc_checks"):
            p(f"- QC: {', '.join(s['qc_checks'])}")
        if s.get("assumed_skills"):
            p(f"- Assumes: {', '.join(s['assumed_skills'])}")
        if parallel:
            p(f"- Can run in parallel with: {parallel}")
        if s.get("notes"):
            p(f"> {s['notes']}")
        p()

    p("## Materials")
    for m in plan.get("materials") or []:
        shelf = f"; shelf life {m.get('shelf_life_days')} days" if m.get("shelf_life_days") else ""
        catalog = f", catalog {m.get('catalog_no')}" if m.get("catalog_no") else ""
        supplier = f", supplier {m.get('supplier')}" if m.get("supplier") else ""
        order = f", order {m.get('order_priority')}" if m.get("order_priority") else ""
        storage = f", storage {m.get('storage')}{shelf}" if m.get("storage") or shelf else ""
        p(
            f"- **{m.get('name','')}**{catalog}{supplier}: "
            f"{m.get('qty','')} {m.get('unit_size','')} at "
            f"${m.get('unit_cost_usd', 0):.2f} each; total ${m.get('total_cost_usd', 0):.2f}"
            f"{order}{storage}."
        )
    p()

    p("## Equipment")
    for e in plan.get("equipment") or []:
        p(f"- **{e.get('name','')}** ({e.get('model','')}) — {e.get('location','')}")
    p()

    b = plan.get("budget") or {}
    p("## Budget")
    p(f"**Total: ${b.get('total_usd', 0):,.0f} {b.get('currency','USD')}** · {b.get('contingency_pct', 0):.0f}% contingency")
    p()
    if b.get("categories"):
        for c in b["categories"]:
            p(f"- {c.get('name')}: ${c.get('total_usd', 0):,.0f}")
        p()
    if plan.get("budget_justification"):
        p(f"> {plan['budget_justification']}")
        p()

    p("## Timeline")
    for ph in plan.get("timeline") or []:
        deps = f"  *(depends on: {', '.join(ph.get('dependencies') or [])})*" if ph.get("dependencies") else ""
        p(f"- **w{ph.get('week_start')}–w{ph.get('week_end')} {ph.get('name','')}**{deps}")
        for d in ph.get("deliverables") or []:
            p(f"  - {d}")
    p()

    p("## Staffing")
    for s in plan.get("staffing") or []:
        p(f"- **{s.get('role','')}** — {s.get('named_person','')} ({s.get('institution','')}) · {s.get('fte_pct',0)}% FTE")
        if s.get("expertise_tags"):
            p(f"  - Expertise: {', '.join(s['expertise_tags'])}")
    p()

    v = plan.get("validation") or {}
    p("## Validation")
    p()
    p("**Success criteria:**")
    for c in v.get("success_criteria") or []:
        p(f"- {c}")
    if v.get("failure_modes"):
        p()
        p("**Failure modes:**")
        for f in v["failure_modes"]:
            p(f"- *{f}*")
    if v.get("statistics_plan"):
        p()
        p(f"**Statistics.** {v['statistics_plan']}")
    p()

    if plan.get("open_questions"):
        p("## Open questions for the scientist")
        for q in plan["open_questions"]:
            p(f"- {q}")
        p()

    if plan.get("references"):
        p("## References")
        for r in plan["references"]:
            line = f"- *{r.get('title','')}* — {r.get('authors','')}"
            if r.get("year"):
                line += f" ({r['year']})"
            if r.get("doi") or r.get("url"):
                line += f"  [{r.get('doi') or r.get('url')}]({r.get('url') or 'https://doi.org/' + (r.get('doi') or '')})"
            p(line)
        p()

    return "\n".join(lines)


# ─── HTML for PDF (xhtml2pdf) ─────────────────────────────────────────────
# Palette mirrors the frontend tokens (src/index.css):
#   --primary       hsl(192 75% 38%)  → #178bb6   (ocean teal)
#   --primary-deep  hsl(192 80% 26%)  → #0c5e80   (deep teal)
#   --foreground    hsl(220 30% 14%)  → #1a2332   (near-navy)
#   --secondary     hsl(40 36% 92%)   → #f0eee8   (warm cream)
#   --accent / brass hsl(28 65% 50%)  → #d28338   (brass)
#   --status-novel  hsl(150 50% 35%)  → #2d8c5f   (kelp green)
PDF_CSS = """
@page { size: Letter; margin: 0.9in 0.95in 1.0in; @bottom-center { content: "aethena · page " counter(page) " of " counter(pages); font-family: Georgia, serif; font-size: 8pt; color: #6e6e6e; } }
body { font-family: 'Fraunces', 'EB Garamond', 'Garamond', Georgia, serif; color: #1a2332; line-height: 1.6; font-size: 10.5pt; }
h1 { font-size: 24pt; margin: 0 0 12pt; color: #0c5e80; font-weight: 600; line-height: 1.15; letter-spacing: -0.01em; border-bottom: 2.5pt solid #178bb6; padding-bottom: 10pt; }
h2 { font-size: 15pt; margin: 22pt 0 8pt; color: #178bb6; font-weight: 600; padding: 4pt 8pt; background: linear-gradient(180deg, #e8f4fa, #f4faff); border-left: 3pt solid #178bb6; letter-spacing: -0.005em; }
h3 { font-size: 11.5pt; margin: 12pt 0 4pt; font-style: italic; color: #2a3a4a; font-weight: 600; }
p { margin: 4pt 0 6pt; }
blockquote { border-left: 3pt solid #d28338; padding: 8pt 12pt; color: #4a3a28; font-style: italic; margin: 10pt 0; background: #fbf6ee; border-radius: 0 4pt 4pt 0; }
code { font-family: 'JetBrains Mono', 'Menlo', monospace; font-size: 9pt; background: #f0eee8; padding: 1pt 4pt; border-radius: 2pt; color: #a06424; }
pre { font-family: 'JetBrains Mono', 'Menlo', monospace; font-size: 8.5pt; background: #f0eee8; padding: 8pt; border-radius: 3pt; border: 0.5pt solid #e0dccf; }
table { width: 100%; border-collapse: collapse; margin: 8pt 0 14pt; font-size: 8.75pt; table-layout: fixed; border: 0.5pt solid #c8d4dc; border-radius: 3pt; overflow: hidden; }
th, td { padding: 6pt 8pt; border-bottom: 0.5pt solid #d8dee4; text-align: left; vertical-align: top; word-wrap: break-word; overflow-wrap: break-word; }
th { background: #0c5e80; color: #ffffff; text-transform: uppercase; letter-spacing: 0.10em; font-size: 7.75pt; font-weight: 600; border-bottom: none; }
tr:nth-child(even) td { background: #f5f9fc; }
ul, ol { padding-left: 20pt; margin: 4pt 0 8pt; }
li { margin: 3pt 0; }
strong { color: #0c5e80; font-weight: 600; }
em { color: #4a5564; }
a { color: #178bb6; text-decoration: none; border-bottom: 0.5pt solid #b8d8e6; }
hr { border: none; border-top: 0.5pt dashed #c8d4dc; margin: 14pt 0; }
p, li, blockquote { page-break-inside: avoid; }
h1, h2 { page-break-after: avoid; }
"""


def markdown_to_html(md_text: str) -> str:
    import markdown as md_mod

    body = md_mod.markdown(md_text, extensions=["tables", "fenced_code"])
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{PDF_CSS}</style></head><body>{body}</body></html>"""


def render_pdf(md_text: str) -> bytes:
    """Use xhtml2pdf — pure Python, no system deps."""
    import io

    from xhtml2pdf import pisa

    html = markdown_to_html(md_text)
    buf = io.BytesIO()
    pisa.CreatePDF(html, dest=buf, encoding="utf-8")
    return buf.getvalue()


def render_docx(md_text: str) -> bytes:
    """Use pandoc if available, else fall back to python-docx (less pretty)."""
    pandoc = _pandoc()
    if pandoc:
        with tempfile.TemporaryDirectory() as td:
            tdp = Path(td)
            (tdp / "in.md").write_text(md_text, encoding="utf-8")
            out = tdp / "out.docx"
            subprocess.run(
                [pandoc, str(tdp / "in.md"), "-o", str(out), "--standalone"],
                check=True,
                capture_output=True,
            )
            return out.read_bytes()
    # fallback
    from docx import Document

    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    buf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(buf.name)
    return Path(buf.name).read_bytes()


def render_latex(md_text: str) -> str:
    pandoc = _pandoc()
    if pandoc:
        with tempfile.TemporaryDirectory() as td:
            tdp = Path(td)
            (tdp / "in.md").write_text(md_text, encoding="utf-8")
            out = tdp / "out.tex"
            subprocess.run(
                [pandoc, str(tdp / "in.md"), "-o", str(out), "--standalone"],
                check=True,
                capture_output=True,
            )
            return out.read_text(encoding="utf-8")
    # fallback: minimal hand-rolled
    import re

    tex = md_text
    tex = re.sub(r"^# (.+)$", r"\\section*{\1}", tex, flags=re.M)
    tex = re.sub(r"^## (.+)$", r"\\subsection*{\1}", tex, flags=re.M)
    tex = re.sub(r"^### (.+)$", r"\\subsubsection*{\1}", tex, flags=re.M)
    tex = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", tex)
    tex = re.sub(r"\*(.+?)\*", r"\\textit{\1}", tex)
    return "\\documentclass[11pt]{article}\n\\usepackage{geometry}\n\\geometry{letterpaper, margin=1in}\n\\begin{document}\n" + tex + "\n\\end{document}\n"
